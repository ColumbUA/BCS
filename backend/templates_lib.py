"""Генератор .docx документів для діловодства роти.

Шаблони — типові зразки на основі загальновідомих практик військового діловодства
(рапорти, накази, акти, журнали, доповіді). Користувач має звіряти з керівними
документами своєї військової частини.
"""
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Helpers ----------
def _set_default_font(doc, name="Times New Roman", size=12):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _h(doc, text, align="center", bold=True, size=14):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    return p


def _p(doc, text="", align="justify", bold=False, italic=False, size=12, indent_first=True):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    if indent_first and align == "justify":
        p.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p


def _signature_block(doc, position, rank, fio, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(f"{position}\n").bold = True
    p.add_run(f"{rank}                                         _____________ {fio}")
    return p


def _date_signature_pair(doc):
    """Дата (зліва) і підпис (справа) у одному рядку."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.cell(0, 0).text = "«___» _________________ 20___ р."
    cell_r = table.cell(0, 1)
    cell_r.text = "________________________________"
    cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("(підпис)").italic = True


def _addressee(doc, lines):
    """Шапка-адресат у правому верхньому куті."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ln in lines:
        r = p.add_run(ln + "\n")
        r.font.size = Pt(12)


def _table_header(table, headers):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True


def _table_widths(table, cm_widths):
    for i, w in enumerate(cm_widths):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)


# ---------- Контекст ----------
def make_context(soldier=None, settings=None, extras=None):
    s = soldier or {}
    st = settings or {}
    ctx = {
        "fio": s.get("fio", "_______________________________"),
        "callsign": s.get("callsign", ""),
        "rank": s.get("rank", "_______________"),
        "position": s.get("position", "_______________________________"),
        "node_path": s.get("node_path", "______________________"),
        "birth_date": s.get("birth_date", "__.__.____"),
        "mobilized_at": s.get("mobilized_at", "__.__.____"),
        "today": datetime.now().strftime("%d.%m.%Y"),
        "year": datetime.now().year,
        # Реквізити частини
        "unit_full": st.get("unit_full", "Військова частина _______________"),
        "unit_short": st.get("unit_short", "В/Ч А____"),
        "unit_chief": st.get("unit_chief", "________________________________"),
        "unit_chief_rank": st.get("unit_chief_rank", "____________"),
        "city": st.get("city", "_______________"),
        "company_name": st.get("company_name", "Рота радіо та радіотехнічної розвідки"),
        "company_chief": st.get("company_chief", "командиру роти"),
    }
    if extras:
        ctx.update(extras)
    return ctx


# ============================ ШАБЛОНИ ============================

def _new_doc():
    doc = Document()
    _set_default_font(doc)
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    return doc


# ---------- РАПОРТИ ----------

def report_vacation(ctx):
    doc = _new_doc()
    _addressee(doc, [f"Командиру {ctx['unit_short']}", ctx["unit_chief_rank"], ctx["unit_chief"]])
    doc.add_paragraph()
    _h(doc, "РАПОРТ", align="center")
    _p(doc, f"Прошу Вашого дозволу надати мені щорічну основну відпустку "
            f"тривалістю 30 (тридцять) календарних днів за {ctx['year']} рік.")
    _p(doc, "Місце проведення відпустки: _________________________________________________.")
    _p(doc, "Контактний номер телефону на період відпустки: ________________________.")
    _p(doc, "За попередні періоди відпусткою не користувався/користувалася частково "
            "(зазначити дати): _____________________________________________.")
    doc.add_paragraph()
    _date_signature_pair(doc)
    doc.add_paragraph()
    _p(doc, f"{ctx['position']}", indent_first=False)
    _p(doc, f"{ctx['rank']}                                                "
            f"{ctx['fio']}", indent_first=False)
    return doc


def report_dismissal_health(ctx):
    doc = _new_doc()
    _addressee(doc, [f"Командиру {ctx['unit_short']}", ctx["unit_chief_rank"], ctx["unit_chief"]])
    doc.add_paragraph()
    _h(doc, "РАПОРТ", align="center")
    _p(doc, "Прошу Вашого направлення мене на військово-лікарську комісію (ВЛК) "
            "для визначення придатності до подальшого проходження військової служби.")
    _p(doc, "Підстава: погіршення стану здоров'я (зазначити захворювання/травму): "
            "____________________________________________________________________________.")
    _p(doc, "Медичні документи (виписка/направлення/довідка) додаються: _____ арк.")
    doc.add_paragraph()
    _date_signature_pair(doc)
    doc.add_paragraph()
    _p(doc, f"{ctx['position']}", indent_first=False)
    _p(doc, f"{ctx['rank']}                                                "
            f"{ctx['fio']}", indent_first=False)
    return doc


def report_financial_aid(ctx):
    doc = _new_doc()
    _addressee(doc, [f"Командиру {ctx['unit_short']}", ctx["unit_chief_rank"], ctx["unit_chief"]])
    doc.add_paragraph()
    _h(doc, "РАПОРТ", align="center")
    _p(doc, f"Прошу Вашого дозволу надати мені одноразову грошову допомогу для вирішення "
            f"соціально-побутових питань відповідно до чинного законодавства.")
    _p(doc, "Підстава: ________________________________________________________________________.")
    _p(doc, "Підтверджуючі документи додаються: ____ арк.")
    doc.add_paragraph()
    _date_signature_pair(doc)
    doc.add_paragraph()
    _p(doc, f"{ctx['position']}", indent_first=False)
    _p(doc, f"{ctx['rank']}                                                "
            f"{ctx['fio']}", indent_first=False)
    return doc


def report_incident(ctx):
    doc = _new_doc()
    _addressee(doc, [f"Командиру {ctx['unit_short']}", ctx["unit_chief_rank"], ctx["unit_chief"]])
    doc.add_paragraph()
    _h(doc, "РАПОРТ", align="center")
    _p(doc, "Доповідаю, що «___» _________ 20___ року приблизно о ______ год. ______ хв. "
            "за адресою/координатами: ____________________________________________________ "
            "сталася подія: _________________________________________________________________.")
    _p(doc, "Обставини події: _____________________________________________________________________"
            "________________________________________________________________________________"
            "________________________________________________________________________________.")
    _p(doc, "Свідки події: ________________________________________________________________________.")
    _p(doc, "Заходи, вжиті мною: __________________________________________________________________"
            "________________________________________________________________________________.")
    doc.add_paragraph()
    _date_signature_pair(doc)
    doc.add_paragraph()
    _p(doc, f"{ctx['position']}", indent_first=False)
    _p(doc, f"{ctx['rank']}                                                "
            f"{ctx['fio']}", indent_first=False)
    return doc


def report_business_trip(ctx):
    doc = _new_doc()
    _addressee(doc, [f"Командиру {ctx['unit_short']}", ctx["unit_chief_rank"], ctx["unit_chief"]])
    doc.add_paragraph()
    _h(doc, "РАПОРТ", align="center")
    _p(doc, "Прошу Вашого дозволу направити мене у службове відрядження.")
    _p(doc, "Місце призначення: ___________________________________________________________________.")
    _p(doc, "Мета відрядження: ____________________________________________________________________.")
    _p(doc, "Період: з «___» _________ 20___ р. по «___» _________ 20___ р. ( ___ діб).")
    _p(doc, "Транспорт: ___________________________________________________________________________.")
    doc.add_paragraph()
    _date_signature_pair(doc)
    doc.add_paragraph()
    _p(doc, f"{ctx['position']}", indent_first=False)
    _p(doc, f"{ctx['rank']}                                                "
            f"{ctx['fio']}", indent_first=False)
    return doc


def report_promotion(ctx):
    doc = _new_doc()
    _addressee(doc, [f"Командиру {ctx['unit_short']}", ctx["unit_chief_rank"], ctx["unit_chief"]])
    doc.add_paragraph()
    _h(doc, "РАПОРТ", align="center")
    _p(doc, "Прошу Вашого клопотання про присвоєння мені чергового військового звання "
            "«____________________».")
    _p(doc, f"Чинне звання: {ctx['rank']}, присвоєне «___» _________ 20___ р.")
    _p(doc, f"Посада, яку обіймаю: {ctx['position']} ({ctx['node_path']}).")
    _p(doc, f"Дата мобілізації / прийняття на службу: {ctx['mobilized_at']}.")
    _p(doc, "Стягнення відсутні / вказати: _________________________________________________________.")
    doc.add_paragraph()
    _date_signature_pair(doc)
    doc.add_paragraph()
    _p(doc, f"{ctx['position']}", indent_first=False)
    _p(doc, f"{ctx['rank']}                                                "
            f"{ctx['fio']}", indent_first=False)
    return doc


def report_dismissal_contract(ctx):
    doc = _new_doc()
    _addressee(doc, [f"Командиру {ctx['unit_short']}", ctx["unit_chief_rank"], ctx["unit_chief"]])
    doc.add_paragraph()
    _h(doc, "РАПОРТ", align="center")
    _p(doc, "Прошу звільнити мене з військової служби у зв'язку із закінченням строку "
            "контракту про проходження військової служби.")
    _p(doc, "Контракт укладено «___» _________ 20___ р., строк дії — до «___» _________ 20___ р.")
    _p(doc, "Місце реєстрації після звільнення: ____________________________________________________.")
    _p(doc, "Прошу провести зі мною повний розрахунок та видати документи у встановленому порядку.")
    doc.add_paragraph()
    _date_signature_pair(doc)
    doc.add_paragraph()
    _p(doc, f"{ctx['position']}", indent_first=False)
    _p(doc, f"{ctx['rank']}                                                "
            f"{ctx['fio']}", indent_first=False)
    return doc


def report_handover(ctx):
    doc = _new_doc()
    _addressee(doc, [f"Командиру {ctx['unit_short']}", ctx["unit_chief_rank"], ctx["unit_chief"]])
    doc.add_paragraph()
    _h(doc, "РАПОРТ", align="center")
    _p(doc, "Доповідаю, що «___» _________ 20___ р. я прийняв (здав) справи та посаду "
            "____________________________________________ від (кому) "
            "_______________________________________________.")
    _p(doc, "Стан особового складу: ________________________________________________________________.")
    _p(doc, "Стан озброєння та техніки: ____________________________________________________________.")
    _p(doc, "Стан документації, печаток, бланків суворого обліку: __________________________________.")
    _p(doc, "Виявлені нестачі/невідповідності: _____________________________________________________.")
    doc.add_paragraph()
    _date_signature_pair(doc)
    doc.add_paragraph()
    _p(doc, f"{ctx['position']}", indent_first=False)
    _p(doc, f"{ctx['rank']}                                                "
            f"{ctx['fio']}", indent_first=False)
    return doc


# ---------- НАКАЗИ / РОЗПОРЯДЖЕННЯ ----------

def order_combat(ctx):
    doc = _new_doc()
    _h(doc, "БОЙОВИЙ НАКАЗ", align="center", size=16)
    _h(doc, f"командира {ctx['company_name'].lower()}",
       align="center", bold=False, size=12)
    _p(doc, f"№ ____", align="center", indent_first=False)
    _p(doc, f"{ctx['city']}                                                                "
            f"«___» _________ 20___ р.", indent_first=False)
    doc.add_paragraph()
    _p(doc, "1. Противник: _______________________________________________________________________.")
    _p(doc, "2. Свої війська (старший начальник): _________________________________________________.")
    _p(doc, "3. Завдання роти: ____________________________________________________________________"
            "________________________________________________________________________________.")
    _p(doc, "4. Задум бойових дій: ________________________________________________________________"
            "________________________________________________________________________________.")
    _p(doc, "5. Завдання підрозділам:", bold=True)
    _p(doc, "5.1. 1 взводу радіорозвідки — _________________________________________________________.")
    _p(doc, "5.2. 2 взводу радіорозвідки — _________________________________________________________.")
    _p(doc, "5.3. Взводу радіоелектронної розвідки — ______________________________________________.")
    _p(doc, "5.4. Взводу безпілотних авіаційних комплексів — ______________________________________.")
    _p(doc, "5.5. Групі обробки інформації — _______________________________________________________.")
    _p(doc, "6. Місця, час і порядок розгортання: _________________________________________________.")
    _p(doc, "7. Взаємодія: ________________________________________________________________________.")
    _p(doc, "8. Забезпечення (тил, медичне, технічне): ____________________________________________.")
    _p(doc, "9. Управління та зв'язок: ____________________________________________________________.")
    _p(doc, "10. Готовність до виконання — «___» год «___» хв «___» __________ 20___ р.")
    doc.add_paragraph()
    _p(doc, f"Командир {ctx['company_name'].lower()}", indent_first=False)
    _p(doc, "_____________________                              _____________ _________________",
       indent_first=False)
    _p(doc, "(військове звання)                                       (підпис)        (П.І.Б.)",
       indent_first=False)
    return doc


def stroyova_zapyska(ctx):
    doc = _new_doc()
    _h(doc, "СТРОЙОВА ЗАПИСКА", align="center", size=16)
    _h(doc, f"{ctx['company_name'].lower()} станом на «___» _________ 20___ р.",
       align="center", bold=False, size=11)
    doc.add_paragraph()
    rows = [
        ["Категорія", "За штатом", "Списково", "В наявності", "Відсутні"],
        ["Офіцерів", "", "", "", ""],
        ["Сержантів", "", "", "", ""],
        ["Солдатів", "", "", "", ""],
        ["Усього", "", "", "", ""],
    ]
    table = doc.add_table(rows=len(rows), cols=5)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.cell(i, j).text = val
            if i == 0:
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs: r.bold = True
    _p(doc, "")
    _p(doc, "Перелік відсутніх (П.І.Б., підстава, дата вибуття, очікуване повернення):", bold=True, indent_first=False)
    for i in range(1, 6):
        _p(doc, f"{i}. ___________________________________________________________________________________.",
           indent_first=False)
    doc.add_paragraph()
    _p(doc, "Старшина роти ____________________ ____________ _________________________",
       indent_first=False)
    _p(doc, "                  (звання)        (підпис)             (П.І.Б.)",
       indent_first=False)
    return doc


def naryad_order(ctx):
    doc = _new_doc()
    _h(doc, "НАРЯД", align="center", size=16)
    _h(doc, f"на ___________________________ «___» _________ 20___ р.",
       align="center", bold=False, size=12)
    doc.add_paragraph()
    headers = ["№", "Звання", "Прізвище та ініціали", "Посада", "Підрозділ", "Підстава", "Підпис"]
    table = doc.add_table(rows=11, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 11):
        table.cell(i, 0).text = str(i)
    doc.add_paragraph()
    _p(doc, f"Командир {ctx['company_name'].lower()}", indent_first=False)
    _p(doc, "_____________________                              _____________ _________________",
       indent_first=False)
    return doc


# ---------- АКТИ ----------

def act_handover(ctx):
    doc = _new_doc()
    _addressee(doc, ["ЗАТВЕРДЖУЮ", f"Командир {ctx['unit_short']}",
                     ctx["unit_chief_rank"], f"_____________ {ctx['unit_chief']}",
                     "«___» _________ 20___ р."])
    doc.add_paragraph()
    _h(doc, "АКТ", align="center", size=16)
    _h(doc, "прийому-передачі майна", align="center", bold=False, size=12)
    _p(doc, f"{ctx['city']}                                                              "
            f"«___» _________ 20___ р.", indent_first=False)
    doc.add_paragraph()
    _p(doc, "Цей акт складено комісією у складі:")
    _p(doc, "Голова комісії: _______________________________________________________________________.")
    _p(doc, "Члени комісії: _______________________________________________________________________ "
            "_____________________________________________________________________________________.")
    _p(doc, "У присутності:")
    _p(doc, "Здав: __________________________________________________________________________________.")
    _p(doc, "Прийняв: _______________________________________________________________________________.")
    _p(doc, "Перевірено наявність та технічний стан майна за переліком:", bold=True)
    headers = ["№", "Найменування", "Од.", "Кількість", "Стан", "Примітки"]
    table = doc.add_table(rows=11, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 11):
        table.cell(i, 0).text = str(i)
    doc.add_paragraph()
    _p(doc, "Висновок комісії: ______________________________________________________________________.")
    doc.add_paragraph()
    _p(doc, "Голова комісії    _____________ ___________________________", indent_first=False)
    _p(doc, "Члени комісії     _____________ ___________________________", indent_first=False)
    _p(doc, "                  _____________ ___________________________", indent_first=False)
    _p(doc, "Здав              _____________ ___________________________", indent_first=False)
    _p(doc, "Прийняв           _____________ ___________________________", indent_first=False)
    return doc


def act_writeoff(ctx):
    doc = _new_doc()
    _addressee(doc, ["ЗАТВЕРДЖУЮ", f"Командир {ctx['unit_short']}",
                     ctx["unit_chief_rank"], f"_____________ {ctx['unit_chief']}",
                     "«___» _________ 20___ р."])
    doc.add_paragraph()
    _h(doc, "АКТ", align="center", size=16)
    _h(doc, "списання матеріальних цінностей", align="center", bold=False, size=12)
    _p(doc, f"№ ____ від «___» _________ 20___ р.", align="center", indent_first=False)
    doc.add_paragraph()
    _p(doc, "Комісія у складі: ______________________________________________________________________ "
            "________________________________________________________________________________________ "
            "склала цей акт про те, що нижчезазначене майно підлягає списанню з обліку у зв'язку з:")
    _p(doc, "□ непридатністю до подальшого використання (фізичне зношення)")
    _p(doc, "□ втратою / пошкодженням під час бойових дій")
    _p(doc, "□ перевищенням строку експлуатації")
    _p(doc, "□ іншою причиною: _____________________________________________________________________.")
    headers = ["№", "Найменування", "Од.", "Кількість", "Балансова вартість", "Причина"]
    table = doc.add_table(rows=11, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 11):
        table.cell(i, 0).text = str(i)
    doc.add_paragraph()
    _p(doc, "Загальна балансова вартість списаного майна: _____________________________ грн.", indent_first=False)
    doc.add_paragraph()
    _p(doc, "Голова комісії     _____________ ___________________________", indent_first=False)
    _p(doc, "Члени комісії      _____________ ___________________________", indent_first=False)
    _p(doc, "                   _____________ ___________________________", indent_first=False)
    _p(doc, "Матеріально-відповідальна особа _____________ ___________________________", indent_first=False)
    return doc


def act_inventory(ctx):
    doc = _new_doc()
    _h(doc, "АКТ ІНВЕНТАРИЗАЦІЇ", align="center", size=16)
    _h(doc, "матеріальних цінностей", align="center", bold=False, size=12)
    _p(doc, f"№ ____ від «___» _________ 20___ р.", align="center", indent_first=False)
    doc.add_paragraph()
    _p(doc, "Підстава: наказ командира в/ч № ____ від «___» _________ 20___ р.")
    _p(doc, "Місце проведення: _____________________________________________________________________.")
    _p(doc, "Період інвентаризації: з «___» _________ по «___» _________ 20___ р.")
    _p(doc, "Матеріально-відповідальна особа: ______________________________________________________.")
    headers = ["№", "Найменування", "Од.", "За обліком", "Фактично", "Розбіжність", "Примітки"]
    table = doc.add_table(rows=12, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 12):
        table.cell(i, 0).text = str(i)
    doc.add_paragraph()
    _p(doc, "Висновок: _______________________________________________________________________________.")
    _p(doc, "Голова комісії     _____________ ___________________________", indent_first=False)
    _p(doc, "Члени комісії      _____________ ___________________________", indent_first=False)
    _p(doc, "                   _____________ ___________________________", indent_first=False)
    return doc


def act_loss(ctx):
    doc = _new_doc()
    _addressee(doc, ["ЗАТВЕРДЖУЮ", f"Командир {ctx['unit_short']}",
                     ctx["unit_chief_rank"], f"_____________ {ctx['unit_chief']}",
                     "«___» _________ 20___ р."])
    doc.add_paragraph()
    _h(doc, "АКТ", align="center", size=16)
    _h(doc, "про втрату (пошкодження) майна", align="center", bold=False, size=12)
    doc.add_paragraph()
    _p(doc, f"{ctx['city']}                                                              "
            f"«___» _________ 20___ р.", indent_first=False)
    doc.add_paragraph()
    _p(doc, "Комісія у складі: ______________________________________________________________________ "
            "склала цей акт про те, що під час виконання бойового завдання (службових обов'язків) "
            "сталася втрата (пошкодження) такого майна:")
    headers = ["№", "Найменування", "Од.", "Кількість", "Інв.№/SN", "Обставини"]
    table = doc.add_table(rows=6, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 6):
        table.cell(i, 0).text = str(i)
    doc.add_paragraph()
    _p(doc, "Обставини, за яких сталася втрата (пошкодження):")
    _p(doc, "________________________________________________________________________________________"
            "________________________________________________________________________________________"
            "________________________________________________________________________________________.")
    _p(doc, "Заходи, вжиті для пошуку (відновлення): _______________________________________________.")
    _p(doc, "Висновок про наявність вини військовослужбовця: _______________________________________.")
    doc.add_paragraph()
    _p(doc, "Голова комісії     _____________ ___________________________", indent_first=False)
    _p(doc, "Члени комісії      _____________ ___________________________", indent_first=False)
    return doc


def act_position_handover(ctx):
    doc = _new_doc()
    _addressee(doc, ["ЗАТВЕРДЖУЮ", f"Командир {ctx['unit_short']}",
                     ctx["unit_chief_rank"], f"_____________ {ctx['unit_chief']}",
                     "«___» _________ 20___ р."])
    doc.add_paragraph()
    _h(doc, "АКТ", align="center", size=16)
    _h(doc, "прийому-передачі справ і посади", align="center", bold=False, size=12)
    doc.add_paragraph()
    _p(doc, "Цей акт складено комісією про те, що згідно наказу командира в/ч № ____ "
            "від «___» _________ 20___ р., в період з «___» _________ по «___» _________ 20___ р. "
            "проведено прийом-передачу справ і посади ___________________________________________.")
    _p(doc, "Здав посаду: __________________________________________________________________________.")
    _p(doc, "Прийняв посаду: _______________________________________________________________________.")
    _p(doc, "У ході перевірки встановлено:", bold=True)
    _p(doc, "1. Стан особового складу: _____________________________________________________________.")
    _p(doc, "2. Стан озброєння, військової техніки та боєприпасів: _________________________________.")
    _p(doc, "3. Стан матеріальних засобів і майна: _________________________________________________.")
    _p(doc, "4. Стан службової документації, печаток, бланків суворого обліку: _____________________.")
    _p(doc, "5. Виявлені нестачі/недоліки: _________________________________________________________.")
    doc.add_paragraph()
    _p(doc, "Здав     _____________ ___________________________", indent_first=False)
    _p(doc, "Прийняв  _____________ ___________________________", indent_first=False)
    _p(doc, "Голова комісії     _____________ ___________________________", indent_first=False)
    _p(doc, "Члени комісії      _____________ ___________________________", indent_first=False)
    return doc


# ---------- ЖУРНАЛИ (форми таблиць) ----------

def journal_personnel(ctx):
    doc = _new_doc()
    _h(doc, "ЖУРНАЛ", align="center", size=16)
    _h(doc, f"обліку особового складу {ctx['company_name'].lower()}",
       align="center", bold=False, size=12)
    _p(doc, f"Розпочато: «___» _________ 20___ р.", align="center", indent_first=False)
    _p(doc, f"Закінчено: «___» _________ 20___ р.", align="center", indent_first=False)
    doc.add_page_break()
    headers = ["№", "Дата", "Прізвище та ініціали", "Звання", "Посада", "Прибув з", "Підстава", "Вибув куди", "Примітки"]
    table = doc.add_table(rows=21, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 21):
        table.cell(i, 0).text = str(i)
    return doc


def journal_orders(ctx):
    doc = _new_doc()
    _h(doc, "ЖУРНАЛ", align="center", size=16)
    _h(doc, f"обліку наказів по {ctx['company_name'].lower()}",
       align="center", bold=False, size=12)
    doc.add_page_break()
    headers = ["№", "Дата", "№ наказу", "Зміст наказу", "Виконавець", "Відмітка про виконання"]
    table = doc.add_table(rows=21, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 21):
        table.cell(i, 0).text = str(i)
    return doc


def journal_weapons(ctx):
    doc = _new_doc()
    _h(doc, "ЖУРНАЛ", align="center", size=16)
    _h(doc, "обліку видачі та прийому зброї і боєприпасів",
       align="center", bold=False, size=12)
    doc.add_page_break()
    headers = ["№", "Дата/час", "П.І.Б.", "Зброя (тип, №)", "БК", "Підпис отримав", "Підпис здав", "Дата/час повернення", "Примітки"]
    table = doc.add_table(rows=16, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 16):
        table.cell(i, 0).text = str(i)
    return doc


def journal_briefings(ctx):
    doc = _new_doc()
    _h(doc, "ЖУРНАЛ", align="center", size=16)
    _h(doc, "проведення інструктажів з безпеки військової служби",
       align="center", bold=False, size=12)
    doc.add_page_break()
    headers = ["№", "Дата", "Тема інструктажу", "П.І.Б. того, хто інструктував", "П.І.Б. того, хто отримав", "Підпис інструктора", "Підпис того, хто отримав"]
    table = doc.add_table(rows=16, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 16):
        table.cell(i, 0).text = str(i)
    return doc


def journal_absent(ctx):
    doc = _new_doc()
    _h(doc, "ЖУРНАЛ", align="center", size=16)
    _h(doc, "обліку відсутніх (відпустка, відрядження, лікарня)",
       align="center", bold=False, size=12)
    doc.add_page_break()
    headers = ["№", "П.І.Б.", "Звання", "Посада", "Підстава", "Дата вибуття", "Дата повернення", "Підстава поверн.", "Примітки"]
    table = doc.add_table(rows=21, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 21):
        table.cell(i, 0).text = str(i)
    return doc


def journal_property(ctx):
    doc = _new_doc()
    _h(doc, "КНИГА", align="center", size=16)
    _h(doc, "обліку матеріальних цінностей",
       align="center", bold=False, size=12)
    doc.add_page_break()
    headers = ["№", "Найменування", "Од.", "Інв.№/SN", "Прихід (дата, к-сть)", "Видача (дата, кому, к-сть)", "Залишок", "Примітки"]
    table = doc.add_table(rows=21, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 21):
        table.cell(i, 0).text = str(i)
    return doc


# ---------- ДОНЕСЕННЯ / ДОПОВІДІ ----------

def daily_report(ctx):
    doc = _new_doc()
    _h(doc, "ДОБОВА ДОПОВІДЬ", align="center", size=16)
    _h(doc, f"командира {ctx['company_name'].lower()}",
       align="center", bold=False, size=12)
    _h(doc, f"станом на «___» год «___» хв «___» _________ 20___ р.",
       align="center", bold=False, size=11)
    doc.add_paragraph()
    _p(doc, "1. ОБСТАНОВКА:", bold=True, indent_first=False)
    _p(doc, "1.1. Противник: _______________________________________________________________________.")
    _p(doc, "1.2. Свої війська: ____________________________________________________________________.")
    _p(doc, "2. ВИКОНАННЯ ЗАВДАНЬ ЗА ДОБУ:", bold=True, indent_first=False)
    _p(doc, "________________________________________________________________________________________"
            "________________________________________________________________________________________.")
    _p(doc, "3. ОСОБОВИЙ СКЛАД:", bold=True, indent_first=False)
    rows = [["Категорія", "За штатом", "Списково", "В наявності", "Втрати (за добу)", "Хворі"]]
    rows += [["Офіцерів",""]*5, ["Сержантів",""]*5, ["Солдатів",""]*5, ["Усього",""]*5]
    table = doc.add_table(rows=len(rows), cols=6)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row + [""]*(6-len(row))):
            table.cell(i, j).text = str(val)
    _p(doc, "")
    _p(doc, "4. ОЗБРОЄННЯ ТА ТЕХНІКА: _____________________________________________________________.")
    _p(doc, "5. ВИТРАТИ БК / ПММ / ЗВ'ЯЗКУ: _______________________________________________________.")
    _p(doc, "6. РЕЗУЛЬТАТИ РОЗВІДКИ: ______________________________________________________________.")
    _p(doc, "7. ПЛАН ЗАДАЧ НА НАСТУПНУ ДОБУ: ______________________________________________________.")
    _p(doc, "8. ПОТРЕБИ В ЗАБЕЗПЕЧЕННІ: ___________________________________________________________.")
    doc.add_paragraph()
    _p(doc, f"Командир {ctx['company_name'].lower()}", indent_first=False)
    _p(doc, "_____________________ _____________ ___________________________", indent_first=False)
    return doc


def losses_report(ctx):
    doc = _new_doc()
    _h(doc, "ДОНЕСЕННЯ ПРО ВТРАТИ", align="center", size=16)
    _h(doc, f"за період з «___» _________ по «___» _________ 20___ р.",
       align="center", bold=False, size=11)
    doc.add_paragraph()
    headers = ["№", "П.І.Б.", "Звання", "Посада", "Категорія втрати*", "Дата/час", "Місце (координати)", "Обставини", "Примітки"]
    table = doc.add_table(rows=11, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    for i in range(1, 11):
        table.cell(i, 0).text = str(i)
    _p(doc, "")
    _p(doc, "* Категорії: загиблий, поранений, контужений, зник безвісти, в полоні", italic=True, indent_first=False)
    doc.add_paragraph()
    _p(doc, f"Командир {ctx['company_name'].lower()}", indent_first=False)
    _p(doc, "_____________________ _____________ ___________________________", indent_first=False)
    return doc


def task_completion_report(ctx):
    doc = _new_doc()
    _h(doc, "ДОПОВІДЬ", align="center", size=16)
    _h(doc, "про виконання бойового завдання",
       align="center", bold=False, size=12)
    doc.add_paragraph()
    _p(doc, "Згідно бойового наказу № ____ від «___» _________ 20___ р. підрозділ виконував "
            "завдання: ________________________________________________________________________________.")
    _p(doc, "Період виконання: з «___» _________ по «___» _________ 20___ р.")
    _p(doc, "Сили і засоби, що залучалися: _________________________________________________________.")
    _p(doc, "Хід виконання завдання: _______________________________________________________________"
            "________________________________________________________________________________________"
            "________________________________________________________________________________________.")
    _p(doc, "Результати:", bold=True)
    _p(doc, "— виявлено об'єктів противника: _______________________________________________________.")
    _p(doc, "— перехоплено сеансів зв'язку: ________________________________________________________.")
    _p(doc, "— ідентифіковано засобів РЕБ/зв'язку противника: ______________________________________.")
    _p(doc, "— інше: ______________________________________________________________________________.")
    _p(doc, "Втрати власні: ________________________________________________________________________.")
    _p(doc, "Витрати: БК — ________, ПММ — ________, інше — ________.")
    _p(doc, "Висновки і пропозиції: ________________________________________________________________.")
    doc.add_paragraph()
    _p(doc, f"Командир {ctx['company_name'].lower()}", indent_first=False)
    _p(doc, "_____________________ _____________ ___________________________", indent_first=False)
    return doc


# ============================ КАТАЛОГ ============================

CATALOG = [
    # id, category, name, builder, description
    # Рапорти
    {"id": "report_vacation",          "category": "Рапорти",   "name": "Рапорт на щорічну відпустку",                    "builder": report_vacation,          "desc": "30 діб основної відпустки"},
    {"id": "report_dismissal_health",  "category": "Рапорти",   "name": "Рапорт на ВЛК (стан здоров'я)",                  "builder": report_dismissal_health,  "desc": "Направлення на ВЛК"},
    {"id": "report_financial_aid",     "category": "Рапорти",   "name": "Рапорт на матеріальну допомогу",                 "builder": report_financial_aid,     "desc": "Одноразова допомога"},
    {"id": "report_incident",          "category": "Рапорти",   "name": "Рапорт по факту події",                          "builder": report_incident,          "desc": "Втрата майна / порушення"},
    {"id": "report_business_trip",     "category": "Рапорти",   "name": "Рапорт на службове відрядження",                 "builder": report_business_trip,     "desc": "З мету та строки"},
    {"id": "report_promotion",         "category": "Рапорти",   "name": "Рапорт на чергове військове звання",             "builder": report_promotion,         "desc": ""},
    {"id": "report_dismissal_contract","category": "Рапорти",   "name": "Рапорт на звільнення (закінчення контракту)",   "builder": report_dismissal_contract, "desc": ""},
    {"id": "report_handover",          "category": "Рапорти",   "name": "Рапорт про прийом-здачу справ і посади",        "builder": report_handover,          "desc": ""},
    # Накази
    {"id": "order_combat",             "category": "Накази",    "name": "Бойовий наказ командира роти",                   "builder": order_combat,             "desc": "Стандартна структура з 10 пунктів"},
    {"id": "stroyova_zapyska",         "category": "Накази",    "name": "Стройова записка",                                "builder": stroyova_zapyska,         "desc": "За штатом / списково / в наявності"},
    {"id": "naryad_order",             "category": "Накази",    "name": "Наряд (на охорону / роботи)",                     "builder": naryad_order,             "desc": "Список призначених"},
    # Акти
    {"id": "act_handover",             "category": "Акти",      "name": "Акт прийому-передачі майна",                      "builder": act_handover,             "desc": "З комісією"},
    {"id": "act_writeoff",             "category": "Акти",      "name": "Акт списання матеріальних цінностей",            "builder": act_writeoff,             "desc": ""},
    {"id": "act_inventory",            "category": "Акти",      "name": "Акт інвентаризації",                              "builder": act_inventory,            "desc": "Перевірка фактичних залишків"},
    {"id": "act_loss",                 "category": "Акти",      "name": "Акт про втрату (пошкодження) майна",              "builder": act_loss,                 "desc": ""},
    {"id": "act_position_handover",    "category": "Акти",      "name": "Акт прийому-передачі справ і посади",            "builder": act_position_handover,    "desc": ""},
    # Журнали
    {"id": "journal_personnel",        "category": "Журнали",   "name": "Журнал обліку особового складу",                  "builder": journal_personnel,        "desc": "Прибуття/вибуття"},
    {"id": "journal_orders",           "category": "Журнали",   "name": "Журнал обліку наказів по роті",                   "builder": journal_orders,           "desc": ""},
    {"id": "journal_weapons",          "category": "Журнали",   "name": "Журнал видачі/прийому зброї та БК",               "builder": journal_weapons,          "desc": ""},
    {"id": "journal_briefings",        "category": "Журнали",   "name": "Журнал інструктажів з БВС",                       "builder": journal_briefings,        "desc": "Безпека військової служби"},
    {"id": "journal_absent",           "category": "Журнали",   "name": "Журнал обліку відсутніх",                         "builder": journal_absent,           "desc": "Відпустка/відрядження/лікарня"},
    {"id": "journal_property",         "category": "Журнали",   "name": "Книга обліку матеріальних цінностей",             "builder": journal_property,         "desc": ""},
    # Донесення
    {"id": "daily_report",             "category": "Донесення", "name": "Добова доповідь",                                  "builder": daily_report,             "desc": "Стан, виконання, потреби"},
    {"id": "losses_report",            "category": "Донесення", "name": "Донесення про втрати",                             "builder": losses_report,            "desc": ""},
    {"id": "task_completion_report",   "category": "Донесення", "name": "Доповідь про виконання бойового завдання",        "builder": task_completion_report,   "desc": ""},
]


def list_templates():
    return [{"id": t["id"], "category": t["category"], "name": t["name"], "desc": t["desc"]} for t in CATALOG]


def render_template(template_id, soldier=None, settings=None, extras=None):
    tpl = next((t for t in CATALOG if t["id"] == template_id), None)
    if not tpl:
        return None, None
    ctx = make_context(soldier=soldier, settings=settings, extras=extras)
    doc = tpl["builder"](ctx)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = tpl["name"].replace("/", "-").replace("\\", "-")
    return buf, f"{safe_name}.docx"
